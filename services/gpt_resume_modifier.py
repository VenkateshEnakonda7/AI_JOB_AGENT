import gspread
from google.oauth2.service_account import Credentials
from config import GOOGLE_SHEETS_CREDENTIALS_FILE, SHEET_NAME, OLD_RESUME_PATH
from docx import Document
from docx2pdf import convert
import openai
import config
import os
import re

# Set OpenAI config to OpenRouter
openai.api_key = config.OPENROUTER_API_KEY
openai.base_url = "https://openrouter.ai/api/v1"
client = openai.OpenAI(
        api_key=config.OPENROUTER_API_KEY,
        base_url="https://openrouter.ai/api/v1"
    )

def load_old_resume():
    document = Document(OLD_RESUME_PATH)
    full_text = []
    for para in document.paragraphs:
        full_text.append(para.text)
    return "\n".join(full_text)

def rewrite_resume(old_resume_text, job_description):
    prompt = f"""
                You are a professional resume writer.
                Here is the OLD RESUME:
                ---
                {old_resume_text}
                ---
                Here are the important JOB DESCRIPTION KEYWORDS that must be smartly included into the resume:
                {job_description}

                - Rewrite the resume professionally.
                - Naturally insert the keywords where appropriate.
                - Improve clarity, formatting, and flow.
                - Keep the tone formal, confident, and achievement-driven.
                - DO NOT invent fake experiences.
                - Match the style of a modern resume.

                Only provide the updated resume text. Do not add extra comments.
            """
   
    response = client.chat.completions.create(
        model="mistralai/mistral-7b-instruct",
        messages=[{"role": "user", "content": prompt}]
    )
    return response.choices[0].message.content

# def save_resume(text, company_name):
#     sanitized_company = re.sub(r'\W+', '_', company_name)
#     docx_path = f"Venkatesh_Enankonda_{sanitized_company}.docx"
#     pdf_path = f"Venkatesh_Enankonda_{sanitized_company}.pdf"

#     document = Document()
#     for paragraph in text.split("\n"):
#         if paragraph.strip():
#             document.add_paragraph(paragraph.strip())
#     document.save(docx_path)
#     convert(docx_path, pdf_path)
#     os.remove(docx_path)  # Remove the .docx file after conversion
#     # Optionally, you can also remove the PDF file if needed
#     print(f"✅ Saved: {pdf_path}")

# def extract_keywords_from_description(description):
#     # Simplified keyword extraction (can improve with NLP later)
#     words = re.findall(r'\b\w{4,}\b', description.lower())
#     return list(set(words))

def generate_resumes_from_sheet(job_description):
    # Setup Google Sheet
    # scope = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"]
    # creds = Credentials.from_service_account_file(GOOGLE_SHEETS_CREDENTIALS_FILE, scopes=scope)
    # client = gspread.authorize(creds)
    # sheet = client.open(SHEET_NAME).sheet1

    # rows = sheet.get_all_records()
    old_resume_text = load_old_resume()
    # if not old_resume_text:
    #     print("❌ Old resume text is empty. Please check the OLD_RESUME_PATH.")
    #     return
    # print(f"🔍 Found {len(rows)} rows in the Google Sheet.")
    # row_number = 1  # Assuming the first row is the header
    # column_number = 9  # Assuming the 9th column is where you want to save the updated resume
    # for row in rows:
    #     description = row.get("Job Description")
    #     company = row.get("Company", "UnknownCompany")
    #     if not description:
    #         print(f"⚠️ Skipping row due to empty job description: {row}")
    #         continue

        # keywords = extract_keywords_from_description(description)
        # try:
    return rewrite_resume(old_resume_text, job_description)
        # save_re = save_resume(updated_resume, company)
        # Save the updated resume to Google Sheets (optional)
        # sheet.update_cell(row_number, column_number, save_re)  # Update the cell with the new resume text  
        # row_number += 1  # Increment row number for the next entry
    # print(f"✅ Resume generated for {company}")
        # except gspread.exceptions.APIError as e:
        #     print(f"❌ Google Sheets API error for company {company}: {e}")
        # except Exception as e:
        #     print(f"❌ Error processing company {company}: {e}")
